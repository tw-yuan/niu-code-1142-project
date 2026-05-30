from __future__ import annotations

import io
from pathlib import Path
from typing import Any

from app.config import get_settings
from app.models import GeneratedFile
from app.services.progress_service import write_event
from app.services.system_setting_service import get_int, KEY_MAX_FILE_SIZE_MB, KEY_MAX_FILES_PER_TASK
from app.tools._helpers import (
    ALLOWED_WRITE_EXTENSIONS,
    ToolContext,
    ToolError,
    academic_integrity_footer_markdown,
)
from app.utils.file_utils import make_task_generated_dir, sanitize_filename


MAX_BLOCKS = 200
MAX_BLOCK_TEXT = 2000
MAX_TABLE_ROWS = 100
MAX_TABLE_COLS = 20
MAX_TEXT_CONTENT_CHARS = 30_000
MAX_XLSX_SHEETS = 10
MAX_XLSX_ROWS = 500
MAX_XLSX_COLS = 30
MAX_FILENAME_LEN = 100


def _validate_filename(filename: Any, allowed_exts: set[str]) -> str:
    if not isinstance(filename, str) or not filename.strip():
        raise ToolError("invalid_argument", "filename 必填")
    clean = sanitize_filename(filename)
    if not clean:
        raise ToolError("invalid_argument", "filename 不合法")
    suffix = Path(clean).suffix.lower()
    if suffix not in allowed_exts:
        raise ToolError(
            "invalid_argument",
            f"filename 副檔名必須是 {sorted(allowed_exts)}，目前 {suffix or '(無)'}",
        )
    if len(clean) > MAX_FILENAME_LEN:
        raise ToolError("invalid_argument", f"filename 不可超過 {MAX_FILENAME_LEN} 字元")
    return clean


def _check_caps(ctx: ToolContext, content_size: int) -> None:
    max_files = get_int(ctx.db, KEY_MAX_FILES_PER_TASK, 8)
    if ctx.generated_count >= max_files:
        raise ToolError(
            "size_limit_exceeded",
            f"已達單任務檔案數上限 {max_files}",
        )
    max_size_mb = get_int(ctx.db, KEY_MAX_FILE_SIZE_MB, get_settings().max_file_size_mb)
    max_bytes = max_size_mb * 1024 * 1024
    if content_size > max_bytes:
        raise ToolError(
            "size_limit_exceeded",
            f"檔案大小 {content_size} 超過上限 {max_size_mb}MB",
        )


def _write_to_disk_and_register(
    ctx: ToolContext,
    filename: str,
    purpose: str | None,
    file_format: str,
    raw_bytes: bytes,
) -> dict[str, Any]:
    _check_caps(ctx, len(raw_bytes))
    settings = get_settings()
    out_dir = make_task_generated_dir(Path(settings.generated_file_dir), ctx.task.id)
    target = out_dir / filename

    target.write_bytes(raw_bytes)
    size = target.stat().st_size

    # If a previous GeneratedFile with same filename exists, mark its row as superseded by
    # creating a new record. We do NOT delete the old DB row; old file is overwritten on disk.
    gf = GeneratedFile(
        task_id=ctx.task.id,
        tool_call_id=None,  # filled in by runtime after dispatch returns
        format=file_format,
        filename=filename,
        purpose=(purpose or "")[:500] if purpose else None,
        file_path=str(target),
        size_bytes=size,
        status="success",
    )
    ctx.db.add(gf)
    ctx.db.flush()

    write_event(
        ctx.db,
        ctx.task.id,
        event_type="agent_write",
        message=f"Agent 寫入 {filename}",
        detail={"format": file_format, "size_bytes": size, "purpose": purpose},
    )

    return {
        "generated_file_id": gf.id,
        "filename": filename,
        "size_bytes": size,
    }


def _validate_blocks(blocks: Any) -> list[dict]:
    if not isinstance(blocks, list) or not blocks:
        raise ToolError("invalid_argument", "blocks 需為非空 list")
    if len(blocks) > MAX_BLOCKS:
        raise ToolError("invalid_argument", f"blocks 數量不可超過 {MAX_BLOCKS}")
    valid: list[dict] = []
    for idx, b in enumerate(blocks):
        if not isinstance(b, dict) or "type" not in b:
            raise ToolError("invalid_argument", f"blocks[{idx}] 需為 dict 且包含 type")
        t = b.get("type")
        if t == "heading":
            text = str(b.get("text", ""))[:MAX_BLOCK_TEXT]
            level = b.get("level", 1)
            try:
                level = int(level)
            except (TypeError, ValueError):
                level = 1
            level = max(1, min(level, 6))
            valid.append({"type": "heading", "level": level, "text": text})
        elif t == "paragraph":
            text = str(b.get("text", ""))[:MAX_BLOCK_TEXT]
            valid.append({"type": "paragraph", "text": text})
        elif t in ("bullet_list", "numbered_list"):
            items_raw = b.get("items") or []
            if not isinstance(items_raw, list):
                raise ToolError("invalid_argument", f"blocks[{idx}].items 需為 list")
            items = [str(i)[:MAX_BLOCK_TEXT] for i in items_raw][:MAX_BLOCKS]
            valid.append({"type": t, "items": items})
        elif t == "table":
            cols = b.get("columns") or []
            rows = b.get("rows") or []
            if not isinstance(cols, list) or not isinstance(rows, list):
                raise ToolError("invalid_argument", f"blocks[{idx}] table columns/rows 需為 list")
            cols = [str(c)[:200] for c in cols][:MAX_TABLE_COLS]
            new_rows = []
            for r in rows[:MAX_TABLE_ROWS]:
                if not isinstance(r, list):
                    continue
                new_rows.append([str(c)[:500] for c in r[:MAX_TABLE_COLS]])
            valid.append({"type": "table", "columns": cols, "rows": new_rows})
        else:
            raise ToolError("invalid_argument", f"blocks[{idx}].type 不支援：{t}")
    return valid


# ---------- write_text_file ----------

def tool_write_text_file(ctx: ToolContext, args: dict) -> dict[str, Any]:
    filename = _validate_filename(args.get("filename"), {".txt", ".md"})
    content = args.get("content")
    purpose = args.get("purpose")
    if not isinstance(content, str):
        raise ToolError("invalid_argument", "content 必填且需為字串")
    if len(content) > MAX_TEXT_CONTENT_CHARS:
        content = content[:MAX_TEXT_CONTENT_CHARS]

    full = content + academic_integrity_footer_markdown()
    file_format = "md" if filename.lower().endswith(".md") else "txt"
    return _write_to_disk_and_register(
        ctx,
        filename=filename,
        purpose=purpose if isinstance(purpose, str) else None,
        file_format=file_format,
        raw_bytes=full.encode("utf-8"),
    )


# ---------- write_docx_file ----------

def _docx_blocks_to_bytes(blocks: list[dict]) -> bytes:
    from docx import Document

    doc = Document()
    for b in blocks:
        t = b["type"]
        if t == "heading":
            doc.add_heading(b["text"], level=b["level"])
        elif t == "paragraph":
            doc.add_paragraph(b["text"])
        elif t == "bullet_list":
            for item in b["items"]:
                doc.add_paragraph(item, style="List Bullet")
        elif t == "numbered_list":
            for item in b["items"]:
                doc.add_paragraph(item, style="List Number")
        elif t == "table":
            cols = b["columns"]
            rows = b["rows"]
            if not cols and not rows:
                continue
            n_cols = max(len(cols), max((len(r) for r in rows), default=0))
            n_rows = (1 if cols else 0) + len(rows)
            if n_rows == 0 or n_cols == 0:
                continue
            tbl = doc.add_table(rows=n_rows, cols=n_cols)
            ridx = 0
            if cols:
                for j, c in enumerate(cols):
                    tbl.rows[ridx].cells[j].text = c
                ridx += 1
            for r in rows:
                for j, c in enumerate(r):
                    tbl.rows[ridx].cells[j].text = c
                ridx += 1

    _append_footer_to_docx(doc)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _append_footer_to_docx(doc) -> None:
    footer_text = academic_integrity_footer_markdown().strip()
    for para in footer_text.splitlines():
        doc.add_paragraph(para)


def tool_write_docx_file(ctx: ToolContext, args: dict) -> dict[str, Any]:
    filename = _validate_filename(args.get("filename"), {".docx"})
    purpose = args.get("purpose")
    blocks = _validate_blocks(args.get("blocks"))
    raw = _docx_blocks_to_bytes(blocks)
    return _write_to_disk_and_register(
        ctx,
        filename=filename,
        purpose=purpose if isinstance(purpose, str) else None,
        file_format="docx",
        raw_bytes=raw,
    )


# ---------- write_pdf_file ----------

def _pdf_blocks_to_bytes(blocks: list[dict]) -> bytes:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import mm
    from reportlab.platypus import (
        ListFlowable,
        ListItem,
        Paragraph,
        SimpleDocTemplate,
        Spacer,
        Table,
        TableStyle,
    )

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=A4,
        leftMargin=18 * mm,
        rightMargin=18 * mm,
        topMargin=18 * mm,
        bottomMargin=18 * mm,
    )
    styles = getSampleStyleSheet()
    heading_styles = [
        ParagraphStyle("H1", parent=styles["Heading1"], fontSize=18, spaceAfter=8),
        ParagraphStyle("H2", parent=styles["Heading2"], fontSize=15, spaceAfter=6),
        ParagraphStyle("H3", parent=styles["Heading3"], fontSize=13, spaceAfter=5),
        ParagraphStyle("H4", parent=styles["Heading4"], fontSize=12, spaceAfter=4),
        ParagraphStyle("H5", parent=styles["Heading4"], fontSize=11, spaceAfter=4),
        ParagraphStyle("H6", parent=styles["Heading4"], fontSize=10, spaceAfter=4),
    ]
    body = ParagraphStyle("Body", parent=styles["BodyText"], fontSize=10.5, leading=15)

    story: list[Any] = []
    for b in blocks:
        t = b["type"]
        if t == "heading":
            level = max(1, min(b["level"], 6))
            story.append(Paragraph(_escape(b["text"]), heading_styles[level - 1]))
        elif t == "paragraph":
            story.append(Paragraph(_escape(b["text"]), body))
            story.append(Spacer(1, 4))
        elif t in ("bullet_list", "numbered_list"):
            items = [ListItem(Paragraph(_escape(i), body)) for i in b["items"]]
            bullet_type = "bullet" if t == "bullet_list" else "1"
            story.append(ListFlowable(items, bulletType=bullet_type))
            story.append(Spacer(1, 4))
        elif t == "table":
            data = [list(b["columns"])] if b["columns"] else []
            data.extend([list(r) for r in b["rows"]])
            if data:
                table = Table(data, hAlign="LEFT")
                table.setStyle(
                    TableStyle(
                        [
                            ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
                            ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                            ("VALIGN", (0, 0), (-1, -1), "TOP"),
                            ("FONTSIZE", (0, 0), (-1, -1), 9),
                        ]
                    )
                )
                story.append(table)
                story.append(Spacer(1, 6))

    story.append(Spacer(1, 12))
    footer_text = academic_integrity_footer_markdown().strip()
    footer_style = ParagraphStyle(
        "Footer", parent=body, fontSize=8.5, textColor=colors.HexColor("#444444"), leading=12,
    )
    for line in footer_text.splitlines():
        if line.strip():
            story.append(Paragraph(_escape(line), footer_style))

    doc.build(story)
    return buf.getvalue()


def _escape(s: str) -> str:
    return (
        s.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
    )


def tool_write_pdf_file(ctx: ToolContext, args: dict) -> dict[str, Any]:
    filename = _validate_filename(args.get("filename"), {".pdf"})
    purpose = args.get("purpose")
    blocks = _validate_blocks(args.get("blocks"))
    raw = _pdf_blocks_to_bytes(blocks)
    return _write_to_disk_and_register(
        ctx,
        filename=filename,
        purpose=purpose if isinstance(purpose, str) else None,
        file_format="pdf",
        raw_bytes=raw,
    )


# ---------- write_xlsx_file ----------

def _xlsx_sheets_to_bytes(sheets: list[dict]) -> bytes:
    from openpyxl import Workbook

    wb = Workbook()
    wb.remove(wb.active)
    for s in sheets:
        title = (s.get("name") or "Sheet")[:30] or "Sheet"
        ws = wb.create_sheet(title=title)
        cols = s.get("columns") or []
        rows = s.get("rows") or []
        if cols:
            ws.append([str(c) for c in cols[:MAX_XLSX_COLS]])
        for r in rows[:MAX_XLSX_ROWS]:
            ws.append([("" if v is None else str(v)) for v in r[:MAX_XLSX_COLS]])

    # Footer sheet for academic integrity
    note_ws = wb.create_sheet(title="AcademicIntegrity")
    note_ws.append(["學術誠信提醒"])
    for line in academic_integrity_footer_markdown().strip().splitlines():
        if line.strip():
            note_ws.append([line])

    if not wb.sheetnames:
        wb.create_sheet(title="Empty")

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def tool_write_xlsx_file(ctx: ToolContext, args: dict) -> dict[str, Any]:
    filename = _validate_filename(args.get("filename"), {".xlsx"})
    purpose = args.get("purpose")
    sheets = args.get("sheets")
    if not isinstance(sheets, list) or not sheets:
        raise ToolError("invalid_argument", "sheets 需為非空 list")
    if len(sheets) > MAX_XLSX_SHEETS:
        raise ToolError("invalid_argument", f"sheets 數量不可超過 {MAX_XLSX_SHEETS}")
    cleaned: list[dict] = []
    for s in sheets:
        if not isinstance(s, dict):
            raise ToolError("invalid_argument", "sheet 需為 dict")
        cleaned.append(
            {
                "name": str(s.get("name") or "Sheet")[:30],
                "columns": list(s.get("columns") or []),
                "rows": list(s.get("rows") or []),
            }
        )
    raw = _xlsx_sheets_to_bytes(cleaned)
    return _write_to_disk_and_register(
        ctx,
        filename=filename,
        purpose=purpose if isinstance(purpose, str) else None,
        file_format="xlsx",
        raw_bytes=raw,
    )
