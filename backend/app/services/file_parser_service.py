from __future__ import annotations

import csv
import io
import json
import logging
from dataclasses import dataclass
from pathlib import Path
from typing import Any

logger = logging.getLogger(__name__)


MAX_PARSED_TEXT_CHARS = 200_000
MAX_TABLE_ROWS = 500
MAX_TABLE_COLS = 50
SUMMARY_CHARS = 300


@dataclass
class ParseResult:
    parse_status: str  # success | failed | skipped
    parsed_text: str | None
    parsed_table_json: Any | None
    summary: str | None
    error_message: str | None


def parse_file(path: Path, file_type: str) -> ParseResult:
    try:
        if file_type == "pdf":
            return _parse_pdf(path)
        if file_type == "docx":
            return _parse_docx(path)
        if file_type in {"txt", "md"}:
            return _parse_text(path)
        if file_type == "xlsx":
            return _parse_xlsx(path)
        if file_type == "csv":
            return _parse_csv(path)
        if file_type in {"png", "jpg", "webp"}:
            return _parse_image(path)
        return ParseResult(
            parse_status="skipped",
            parsed_text=None,
            parsed_table_json=None,
            summary=f"未支援的格式 ({file_type})，僅保留 metadata。",
            error_message=None,
        )
    except Exception as exc:  # noqa: BLE001
        logger.exception("Parse failed for %s", path)
        return ParseResult(
            parse_status="failed",
            parsed_text=None,
            parsed_table_json=None,
            summary=None,
            error_message=f"{type(exc).__name__}: {exc}"[:500],
        )


def _truncate_text(text: str) -> tuple[str, str]:
    text = text.strip()
    truncated = text[:MAX_PARSED_TEXT_CHARS]
    summary = text[:SUMMARY_CHARS]
    return truncated, summary


def _parse_pdf(path: Path) -> ParseResult:
    import pymupdf  # type: ignore[import-not-found]

    pieces: list[str] = []
    with pymupdf.open(str(path)) as doc:
        for page in doc:
            pieces.append(page.get_text("text"))
    text = "\n".join(pieces)
    parsed_text, summary = _truncate_text(text)
    return ParseResult(
        parse_status="success",
        parsed_text=parsed_text,
        parsed_table_json=None,
        summary=summary,
        error_message=None,
    )


def _parse_docx(path: Path) -> ParseResult:
    from docx import Document

    doc = Document(str(path))
    pieces: list[str] = []
    for para in doc.paragraphs:
        if para.text:
            pieces.append(para.text)
    tables_data: list[dict] = []
    for idx, table in enumerate(doc.tables):
        rows: list[list[str]] = []
        for row in table.rows:
            rows.append([cell.text for cell in row.cells])
            if len(rows) >= MAX_TABLE_ROWS:
                break
        tables_data.append({"index": idx, "rows": rows})
        pieces.append("\n[表格]")
        for row in rows[:20]:
            pieces.append(" | ".join(row))
    text = "\n".join(pieces)
    parsed_text, summary = _truncate_text(text)
    return ParseResult(
        parse_status="success",
        parsed_text=parsed_text,
        parsed_table_json=tables_data or None,
        summary=summary,
        error_message=None,
    )


def _parse_text(path: Path) -> ParseResult:
    raw_bytes = path.read_bytes()
    text: str | None = None
    for enc in ("utf-8", "utf-8-sig", "big5", "cp950", "gb18030", "latin-1"):
        try:
            text = raw_bytes.decode(enc)
            break
        except UnicodeDecodeError:
            continue
    if text is None:
        text = raw_bytes.decode("utf-8", errors="replace")
    parsed_text, summary = _truncate_text(text)
    return ParseResult(
        parse_status="success",
        parsed_text=parsed_text,
        parsed_table_json=None,
        summary=summary,
        error_message=None,
    )


def _parse_xlsx(path: Path) -> ParseResult:
    from openpyxl import load_workbook

    wb = load_workbook(str(path), read_only=True, data_only=True)
    sheets_out: list[dict] = []
    text_pieces: list[str] = []

    for sheet in wb.worksheets:
        rows: list[list[str]] = []
        columns: list[str] = []
        row_count = 0
        for r_idx, row in enumerate(sheet.iter_rows(values_only=True)):
            if r_idx >= MAX_TABLE_ROWS:
                break
            row_count += 1
            cells = [
                "" if cell is None else str(cell)
                for cell in row[:MAX_TABLE_COLS]
            ]
            if r_idx == 0:
                columns = cells
            else:
                rows.append(cells)

        truncated = sheet.max_row is not None and sheet.max_row > MAX_TABLE_ROWS
        sheets_out.append(
            {
                "name": sheet.title,
                "columns": columns,
                "rows": rows,
                "row_count": row_count,
                "truncated": truncated,
            }
        )

        text_pieces.append(f"## Sheet: {sheet.title}")
        text_pieces.append(" | ".join(columns))
        for row in rows[:10]:
            text_pieces.append(" | ".join(row))

    wb.close()
    text = "\n".join(text_pieces)
    parsed_text, summary = _truncate_text(text)
    return ParseResult(
        parse_status="success",
        parsed_text=parsed_text,
        parsed_table_json={"sheets": sheets_out},
        summary=summary,
        error_message=None,
    )


def _parse_csv(path: Path) -> ParseResult:
    raw_bytes = path.read_bytes()
    text_str: str | None = None
    for enc in ("utf-8-sig", "utf-8", "big5", "cp950", "gb18030", "latin-1"):
        try:
            text_str = raw_bytes.decode(enc)
            break
        except UnicodeDecodeError:
            continue
    if text_str is None:
        text_str = raw_bytes.decode("utf-8", errors="replace")

    reader = csv.reader(io.StringIO(text_str))
    rows: list[list[str]] = []
    for r_idx, row in enumerate(reader):
        if r_idx >= MAX_TABLE_ROWS:
            break
        rows.append([str(c) for c in row[:MAX_TABLE_COLS]])

    columns = rows[0] if rows else []
    body = rows[1:] if rows else []
    text_pieces = [" | ".join(columns)]
    for row in body[:10]:
        text_pieces.append(" | ".join(row))
    text = "\n".join(text_pieces)
    parsed_text, summary = _truncate_text(text)
    return ParseResult(
        parse_status="success",
        parsed_text=parsed_text,
        parsed_table_json={
            "sheets": [
                {
                    "name": "csv",
                    "columns": columns,
                    "rows": body,
                    "row_count": len(rows),
                    "truncated": False,
                }
            ]
        },
        summary=summary,
        error_message=None,
    )


def _parse_image(path: Path) -> ParseResult:
    size = path.stat().st_size
    summary = (
        f"圖片檔 ({path.suffix.lower().lstrip('.')})，大小 {size} bytes。"
        "目前僅保留 metadata，Agent 視需要可參考檔名與大小。"
    )
    return ParseResult(
        parse_status="success",
        parsed_text=None,
        parsed_table_json=None,
        summary=summary,
        error_message=None,
    )


def safe_jsonable(value: Any) -> Any:
    """Ensure value can be persisted as JSON (called before DB write)."""
    if value is None:
        return None
    try:
        json.dumps(value, ensure_ascii=False)
        return value
    except (TypeError, ValueError):
        return None
