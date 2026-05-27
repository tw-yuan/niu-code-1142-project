import json
from pathlib import Path
from datetime import datetime

from app.utils.file_utils import get_generated_path, get_named_generated_path


ACADEMIC_NOTICE = "此內容由 AI 輔助生成，僅供參考與學習用途。使用者必須自行檢查、修改並確認內容的正確性，不應直接提交作為作業。所有引用來源需自行驗證。"


def _deliverable_text(deliverable: dict) -> str:
    content = deliverable.get("content", "")
    if isinstance(content, str):
        return content
    return json.dumps(content, ensure_ascii=False, indent=2)


def _deliverable_title(deliverable: dict) -> str:
    return str(deliverable.get("title") or deliverable.get("filename") or "AI 產生檔案")


def _deliverable_filename(deliverable: dict, fmt: str) -> str:
    filename = str(deliverable.get("filename") or _deliverable_title(deliverable))
    if not filename.lower().endswith(f".{fmt}"):
        filename = f"{Path(filename).stem}.{fmt}"
    return filename


def export_deliverable(task_id: str, deliverable: dict) -> str:
    fmt = str(deliverable.get("format", "")).lower().lstrip(".")
    if fmt == "txt":
        return export_deliverable_txt(task_id, deliverable)
    if fmt == "docx":
        return export_deliverable_docx(task_id, deliverable)
    if fmt == "pdf":
        return export_deliverable_pdf(task_id, deliverable)
    if fmt == "xlsx":
        return export_deliverable_xlsx(task_id, deliverable)
    raise ValueError(f"不支援的輸出格式：{fmt}")


def export_deliverable_txt(task_id: str, deliverable: dict) -> str:
    fmt = "txt"
    path = get_named_generated_path(task_id, _deliverable_filename(deliverable, fmt), fmt)
    path.write_text(_deliverable_text(deliverable), encoding="utf-8")
    return str(path)


def export_deliverable_docx(task_id: str, deliverable: dict) -> str:
    from docx import Document
    from docx.shared import Pt

    fmt = "docx"
    path = get_named_generated_path(task_id, _deliverable_filename(deliverable, fmt), fmt)
    doc = Document()

    style = doc.styles["Normal"]
    style.font.size = Pt(12)
    style.font.name = "Arial"

    doc.add_heading(_deliverable_title(deliverable), level=0)
    for paragraph in _deliverable_text(deliverable).split("\n"):
        if paragraph.strip():
            doc.add_paragraph(paragraph)

    doc.save(str(path))
    return str(path)


def export_deliverable_pdf(task_id: str, deliverable: dict) -> str:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import mm
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    fmt = "pdf"
    path = get_named_generated_path(task_id, _deliverable_filename(deliverable, fmt), fmt)

    font_registered = False
    for font_path in [
        "/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttc",
        "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
        "/usr/share/fonts/noto-cjk/NotoSansCJK-Regular.ttc",
    ]:
        if Path(font_path).exists():
            try:
                pdfmetrics.registerFont(TTFont("NotoSansCJK", font_path, subfontIndex=0))
                font_registered = True
                break
            except Exception:
                continue

    font_name = "NotoSansCJK" if font_registered else "Helvetica"
    doc = SimpleDocTemplate(
        str(path),
        pagesize=A4,
        leftMargin=20 * mm,
        rightMargin=20 * mm,
        topMargin=20 * mm,
        bottomMargin=20 * mm,
    )

    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name="CJKTitle", fontName=font_name, fontSize=18, leading=24, spaceAfter=12))
    styles.add(ParagraphStyle(name="CJKBody", fontName=font_name, fontSize=11, leading=16, spaceAfter=6))

    story = [Paragraph(_deliverable_title(deliverable), styles["CJKTitle"]), Spacer(1, 4 * mm)]
    for paragraph in _deliverable_text(deliverable).split("\n"):
        if paragraph.strip():
            story.append(Paragraph(paragraph, styles["CJKBody"]))

    doc.build(story)
    return str(path)


def export_deliverable_xlsx(task_id: str, deliverable: dict) -> str:
    from openpyxl import Workbook

    fmt = "xlsx"
    path = get_named_generated_path(task_id, _deliverable_filename(deliverable, fmt), fmt)
    wb = Workbook()
    ws = wb.active
    ws.title = "Content"

    content = deliverable.get("content", "")
    if isinstance(content, list):
        for row in content:
            if isinstance(row, list):
                ws.append(row)
            elif isinstance(row, dict):
                ws.append(list(row.values()))
            else:
                ws.append([str(row)])
    elif isinstance(content, dict):
        for key, value in content.items():
            ws.append([key, json.dumps(value, ensure_ascii=False) if isinstance(value, (dict, list)) else value])
    else:
        for line in str(content).split("\n"):
            ws.append([line])

    ws.column_dimensions["A"].width = 100
    wb.save(str(path))
    return str(path)


def export_txt(task_id: str, structured: dict) -> str:
    path = get_generated_path(task_id, "txt")
    lines = []
    lines.append(f"# {structured.get('title', 'AI 輔助結果')}")
    lines.append(f"\n生成時間：{datetime.now().strftime('%Y-%m-%d %H:%M')}")
    lines.append(f"\n## 作業需求摘要\n{structured.get('assignment_summary', '')}")

    if structured.get("requirements_breakdown"):
        lines.append("\n## 需求拆解")
        for i, item in enumerate(structured["requirements_breakdown"], 1):
            lines.append(f"{i}. {item}")

    if structured.get("answer_outline"):
        lines.append("\n## 回答大綱")
        for i, item in enumerate(structured["answer_outline"], 1):
            lines.append(f"{i}. {item}")

    lines.append(f"\n## 生成草稿\n{structured.get('generated_draft', '')}")

    if structured.get("references"):
        lines.append("\n## 引用來源")
        for ref in structured["references"]:
            lines.append(f"- {ref.get('source_name', '未知來源')}：{ref.get('quote_or_summary', '')}（用於：{ref.get('used_for', '')}）")

    if structured.get("limitations"):
        lines.append("\n## 限制說明")
        for item in structured["limitations"]:
            lines.append(f"- {item}")

    lines.append(f"\n## 學術誠信提醒\n{structured.get('academic_integrity_notice', ACADEMIC_NOTICE)}")

    if structured.get("human_review_checklist"):
        lines.append("\n## 人工確認清單")
        for item in structured["human_review_checklist"]:
            lines.append(f"☐ {item}")

    content = "\n".join(lines)
    path.write_text(content, encoding="utf-8")
    return str(path)


def export_docx(task_id: str, structured: dict) -> str:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

    path = get_generated_path(task_id, "docx")
    doc = Document()

    style = doc.styles["Normal"]
    style.font.size = Pt(12)
    style.font.name = "Arial"

    doc.add_heading(structured.get("title", "AI 輔助結果"), level=0)
    doc.add_paragraph(f"生成時間：{datetime.now().strftime('%Y-%m-%d %H:%M')}")

    doc.add_heading("作業需求摘要", level=1)
    doc.add_paragraph(structured.get("assignment_summary", ""))

    if structured.get("requirements_breakdown"):
        doc.add_heading("需求拆解", level=1)
        for item in structured["requirements_breakdown"]:
            doc.add_paragraph(item, style="List Bullet")

    if structured.get("answer_outline"):
        doc.add_heading("回答大綱", level=1)
        for item in structured["answer_outline"]:
            doc.add_paragraph(item, style="List Number")

    doc.add_heading("生成草稿", level=1)
    for paragraph in structured.get("generated_draft", "").split("\n"):
        if paragraph.strip():
            doc.add_paragraph(paragraph)

    if structured.get("references"):
        doc.add_heading("引用來源", level=1)
        for ref in structured["references"]:
            doc.add_paragraph(
                f"{ref.get('source_name', '未知來源')}：{ref.get('quote_or_summary', '')}（用於：{ref.get('used_for', '')}）",
                style="List Bullet",
            )

    if structured.get("limitations"):
        doc.add_heading("限制說明", level=1)
        for item in structured["limitations"]:
            doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("學術誠信提醒", level=1)
    p = doc.add_paragraph(structured.get("academic_integrity_notice", ACADEMIC_NOTICE))
    for run in p.runs:
        run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

    if structured.get("human_review_checklist"):
        doc.add_heading("人工確認清單", level=1)
        for item in structured["human_review_checklist"]:
            doc.add_paragraph(f"☐ {item}")

    doc.save(str(path))
    return str(path)


def export_pdf(task_id: str, structured: dict) -> str:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import mm
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, ListFlowable, ListItem
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.lib.colors import HexColor

    path = get_generated_path(task_id, "pdf")

    font_registered = False
    for font_path in ["/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttc",
                       "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
                       "/usr/share/fonts/noto-cjk/NotoSansCJK-Regular.ttc"]:
        if Path(font_path).exists():
            try:
                pdfmetrics.registerFont(TTFont("NotoSansCJK", font_path, subfontIndex=0))
                font_registered = True
                break
            except Exception:
                continue

    font_name = "NotoSansCJK" if font_registered else "Helvetica"

    doc = SimpleDocTemplate(str(path), pagesize=A4,
                            leftMargin=20*mm, rightMargin=20*mm,
                            topMargin=20*mm, bottomMargin=20*mm)

    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name="CJKTitle", fontName=font_name, fontSize=18, leading=24, spaceAfter=12))
    styles.add(ParagraphStyle(name="CJKHeading", fontName=font_name, fontSize=14, leading=18, spaceAfter=8, spaceBefore=12))
    styles.add(ParagraphStyle(name="CJKBody", fontName=font_name, fontSize=11, leading=16, spaceAfter=6))
    styles.add(ParagraphStyle(name="CJKWarning", fontName=font_name, fontSize=11, leading=16, spaceAfter=6, textColor=HexColor("#CC0000")))

    story = []
    story.append(Paragraph(structured.get("title", "AI 輔助結果"), styles["CJKTitle"]))
    story.append(Paragraph(f"生成時間：{datetime.now().strftime('%Y-%m-%d %H:%M')}", styles["CJKBody"]))
    story.append(Spacer(1, 6*mm))

    story.append(Paragraph("作業需求摘要", styles["CJKHeading"]))
    story.append(Paragraph(structured.get("assignment_summary", ""), styles["CJKBody"]))

    if structured.get("requirements_breakdown"):
        story.append(Paragraph("需求拆解", styles["CJKHeading"]))
        for item in structured["requirements_breakdown"]:
            story.append(Paragraph(f"• {item}", styles["CJKBody"]))

    if structured.get("answer_outline"):
        story.append(Paragraph("回答大綱", styles["CJKHeading"]))
        for i, item in enumerate(structured["answer_outline"], 1):
            story.append(Paragraph(f"{i}. {item}", styles["CJKBody"]))

    story.append(Paragraph("生成草稿", styles["CJKHeading"]))
    for para in structured.get("generated_draft", "").split("\n"):
        if para.strip():
            story.append(Paragraph(para, styles["CJKBody"]))

    if structured.get("references"):
        story.append(Paragraph("引用來源", styles["CJKHeading"]))
        for ref in structured["references"]:
            story.append(Paragraph(
                f"• {ref.get('source_name', '未知來源')}：{ref.get('quote_or_summary', '')}",
                styles["CJKBody"],
            ))

    if structured.get("limitations"):
        story.append(Paragraph("限制說明", styles["CJKHeading"]))
        for item in structured["limitations"]:
            story.append(Paragraph(f"• {item}", styles["CJKBody"]))

    story.append(Paragraph("學術誠信提醒", styles["CJKHeading"]))
    story.append(Paragraph(structured.get("academic_integrity_notice", ACADEMIC_NOTICE), styles["CJKWarning"]))

    if structured.get("human_review_checklist"):
        story.append(Paragraph("人工確認清單", styles["CJKHeading"]))
        for item in structured["human_review_checklist"]:
            story.append(Paragraph(f"☐ {item}", styles["CJKBody"]))

    doc.build(story)
    return str(path)


def export_xlsx(task_id: str, structured: dict) -> str:
    from openpyxl import Workbook
    from openpyxl.styles import Font, Alignment

    path = get_generated_path(task_id, "xlsx")
    wb = Workbook()

    # Summary sheet
    ws = wb.active
    ws.title = "Summary"
    ws.append(["標題", structured.get("title", "")])
    ws.append(["生成時間", datetime.now().strftime("%Y-%m-%d %H:%M")])
    ws.append(["作業需求摘要", structured.get("assignment_summary", "")])
    ws.append([])
    ws.append(["需求拆解"])
    for item in structured.get("requirements_breakdown", []):
        ws.append(["", item])
    ws.append([])
    ws.append(["回答大綱"])
    for item in structured.get("answer_outline", []):
        ws.append(["", item])
    ws.column_dimensions["A"].width = 20
    ws.column_dimensions["B"].width = 80

    # Answer sheet
    ws2 = wb.create_sheet("Answer")
    ws2.append(["生成草稿"])
    for line in structured.get("generated_draft", "").split("\n"):
        ws2.append([line])
    ws2.column_dimensions["A"].width = 100

    # References sheet
    ws3 = wb.create_sheet("References")
    ws3.append(["來源名稱", "引用或摘要", "用於"])
    for ref in structured.get("references", []):
        ws3.append([ref.get("source_name", ""), ref.get("quote_or_summary", ""), ref.get("used_for", "")])
    ws3.column_dimensions["A"].width = 30
    ws3.column_dimensions["B"].width = 50
    ws3.column_dimensions["C"].width = 30

    # Checklist sheet
    ws4 = wb.create_sheet("Checklist")
    ws4.append(["學術誠信提醒"])
    ws4.append([structured.get("academic_integrity_notice", ACADEMIC_NOTICE)])
    ws4.append([])
    ws4.append(["人工確認清單"])
    for item in structured.get("human_review_checklist", []):
        ws4.append([f"☐ {item}"])
    ws4.append([])
    ws4.append(["限制說明"])
    for item in structured.get("limitations", []):
        ws4.append([f"• {item}"])
    ws4.column_dimensions["A"].width = 80

    wb.save(str(path))
    return str(path)
